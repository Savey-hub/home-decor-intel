# -*- coding: utf-8 -*-
"""
Generate two Word versions of the 家装建材家具行业竞争情报月报:
  1) 汇报稿 (brief, 8-12 pages)  -> report_brief.docx
  2) 详报   (detail, 20-30 pages) -> report_detail.docx
Every table is followed by a natural-language description
(structure ratio + YoY direction + differentiation + action implication).
Data rule: every figure carries source; unverifiable = 待核实; no fabrication.
"""
import json, os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
D = os.path.join(ROOT, "data")

def load(p):
    with open(os.path.join(D, p), encoding="utf-8") as f:
        return json.load(f)

macro    = load("macro_realestate.json")
plat     = load("platform_dynamics.json")
policy   = load("industry_policy.json")
mh       = load("v2/monthly_highlights.json")
dd       = load("sources/layerC_doudian_compass.json")
srcidx   = load("v2/data_sources_index.json")

# ---------- CJK typography helpers ----------
EASTASIA = "微软雅黑"
EASTASIA_BODY = "宋体"

def set_cell_font(cell, size=10, bold=False, ea=EASTASIA_BODY, color=None, align=None):
    for p in cell.paragraphs:
        if align is not None:
            p.alignment = align
        if not p.runs:
            p.add_run("")
        for r in p.runs:
            r.font.size = Pt(size)
            r.font.bold = bold
            r.font.name = "Arial"
            r._element.rPr.rFonts.set(qn('w:eastAsia'), ea)
            if color:
                r.font.color.rgb = color

def style_doc_defaults(doc):
    st = doc.styles['Normal']
    st.font.name = "Arial"
    st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), EASTASIA_BODY)
    # 1.5 line spacing default
    pf = st.paragraph_format
    pf.line_spacing = 1.5

def _set_run_ea(run, ea):
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:eastAsia'), ea)

def add_title(doc, text, size=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    _set_run_ea(r, EASTASIA)
    return p

def add_subtitle(doc, text, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _set_run_ea(r, EASTASIA_BODY)
    return p

def h1(doc, text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.size = Pt(16); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    _set_run_ea(r, EASTASIA)
    return p

def h2(doc, text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.size = Pt(13.5); r.font.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x5B, 0x9A)
    _set_run_ea(r, EASTASIA)
    return p

def h3(doc, text):
    p = doc.add_heading(level=3)
    r = p.add_run(text)
    r.font.size = Pt(12); r.font.bold = True
    r.font.color.rgb = RGBColor(0x3F, 0x3F, 0x3F)
    _set_run_ea(r, EASTASIA)
    return p

def para(doc, text, size=11, indent=True, color=None, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(size*2)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = bold
    if color: r.font.color.rgb = color
    _set_run_ea(r, EASTASIA_BODY)
    return p

def desc(doc, text):
    """Natural-language interpretation paragraph after a table (shaded)."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run("解读：")
    r.font.size = Pt(10.5); r.font.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x5B, 0x9A)
    _set_run_ea(r, EASTASIA)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5); r2.font.color.rgb = RGBColor(0x33,0x33,0x33)
    _set_run_ea(r2, EASTASIA_BODY)
    # light shading
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),'F2F6FB')
    p._p.get_or_add_pPr().append(shd)
    return p

def bullet(doc, text, size=11):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.left_indent = Pt(18)
    r = p.add_run("• " + text)
    r.font.size = Pt(size)
    _set_run_ea(r, EASTASIA_BODY)
    return p

def shade_row(row, fill="1F3864"):
    for c in row.cells:
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),fill)
        c._tc.get_or_add_tcPr().append(shd)

def make_table(doc, headers, rows, widths=None, header_fill="1F3864",
               header_color=RGBColor(0xFF,0xFF,0xFF), fsize=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, htext in enumerate(headers):
        hdr.cells[i].text = htext
        set_cell_font(hdr.cells[i], size=fsize, bold=True, ea=EASTASIA,
                      color=header_color, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_row(hdr, header_fill)
    for r_i, rowdata in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(rowdata):
            cells[i].text = str(val)
            set_cell_font(cells[i], size=fsize, bold=False, ea=EASTASIA_BODY,
                          align=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT)
        if r_i % 2 == 1:
            for c in cells:
                shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),'F4F6F9')
                c._tc.get_or_add_tcPr().append(shd)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return t

def add_source_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(8.5); r.font.italic = True
    r.font.color.rgb = RGBColor(0x88,0x88,0x88)
    _set_run_ea(r, EASTASIA_BODY)
    return p

def page_break(doc):
    doc.add_page_break()

# =========================================================
#  SHARED CONTENT PIECES
# =========================================================
ASOF = macro.get("asOf","2026-07-16")
WINDOW = "近30天（2026-06-17 至 2026-07-16）"

def yoy_word(v):
    """translate a numeric yoy string into Chinese direction words."""
    if v is None or v == "" or v == "待核实":
        return "待核实"
    import re
    m = re.search(r'-?\d+\.?\d*', str(v).replace(',',''))
    if not m: return str(v)
    n = float(m.group())
    if n <= -10: return f"大幅下滑（{v}）"
    if n <= -5:  return f"明显下滑（{v}）"
    if n < -0.5: return f"小幅下滑（{v}）"
    if n <= 0.5: return f"基本持平（{v}）"
    if n < 5:    return f"小幅上涨（{v}）"
    if n < 10:   return f"明显上涨（{v}）"
    return f"大幅上涨（{v}）"

# ---- cover ----
def cover(doc, subtitle):
    for _ in range(3): doc.add_paragraph()
    add_title(doc, "家装建材家具行业", 26)
    add_title(doc, "竞争情报月报", 26)
    doc.add_paragraph()
    add_subtitle(doc, subtitle, 14)
    doc.add_paragraph()
    add_subtitle(doc, f"数据窗口：{WINDOW} ＋ 本月要闻（自然月）", 11)
    add_subtitle(doc, f"数据截止：{ASOF}", 11)
    add_subtitle(doc, "覆盖子行业：家具 · 装修 · 建材 · 卫浴厨房 · 全屋智能 · 全屋定制 · 灯具光源 · 电工五金", 10)
    doc.add_paragraph(); doc.add_paragraph()
    add_subtitle(doc, "数据来源：国家统计局 · 中指研究院 · 克而瑞 · CBMF/CBMCA/CBDA · 抖店罗盘 · 各上市公司公告 · 权威财经媒体", 9)
    add_subtitle(doc, "编制原则：每条数据标注来源与发布日期，不可核实项标『待核实』，冲突信息双列保留，严禁编造", 9)
    page_break(doc)

# ---- monthly highlights ----
def section_monthly(doc, detail=False):
    h1(doc, "一、本月要闻（自然月：2026-07-01 至今）")
    para(doc, mh["intro"], indent=True)
    para(doc, mh["monthlySummary"], indent=True, bold=False)
    groups = [("宏观数据","macro"),("平台大事","platform"),("政策标准","policy"),("头部商家","merchant")]
    for label, key in groups:
        items = mh["highlights"][key]
        h2(doc, f"1.{'一二三四'[[k for _,k in groups].index(key)]}　{label}（{len(items)}条）")
        rows = [[it["date"], it["title"], it["impact"], it["source"]] for it in items]
        make_table(doc, ["日期","要闻标题","影响","来源"], rows, widths=[2.2,9.5,1.3,3.2])
        if detail:
            for it in items:
                para(doc, f"【{it['date']}｜{it['cat']}｜影响{it['impact']}】{it['title']}", bold=True, size=10.5, indent=False)
                para(doc, it["detail"], size=10.5)
                add_source_line(doc, f"来源：{it['source']}　{it['url']}")
        # description
        hi = sum(1 for it in items if it["impact"]=="高")
        desc(doc, f"本组共{len(items)}条，其中高影响{hi}条。" + _month_group_desc(key))
    page_break(doc)

def _month_group_desc(key):
    if key=="macro":
        return ("三条宏观信号共同确认需求端全面深跌：家具类零售-6.6%、建材类-10.5%均大幅跑输社零大盘+1.0%，"
                "叠加6月BHI环比-16.34进入传统淡季。对天猫家装运营而言，7月大盘承压是既定基调，"
                "增长只能来自结构性机会（品类升级、以旧换新、头部集中），不能指望大盘自然回暖。")
    if key=="platform":
        return ("平台端呈现『冰火两重天』：京东家装以整装+140%逆势高增、国补第三批625亿续力，是少数结构性亮点；"
                "抖音则收紧商家资质合规。天猫应重点关注京东在整装/局改赛道的抢跑，以及国补品类（卫浴/家电）的竞品承接节奏。")
    if key=="policy":
        return ("两大标准同步推高质量与合规门槛：智能卫浴可靠性国标与住宅全案整装服务标准均利好研发/服务投入高的头部品牌，"
                "加速中小品牌与腰部装企出清。天猫招商与流量分配应向合规头部倾斜，规避标准落地后的合规风险商家。")
    if key=="merchant":
        return ("头部业绩集中探底：定制家居三巨头（欧派预降50-60%、志邦由盈转亏、索菲亚预降78-85%）印证行业寒冬，"
                "而红星美凯龙扭亏、兔宝宝获国家科技奖显示结构分化。运营需警惕定制家具赛道整体承压对天猫GMV的传导。")
    return ""

# ---- macro ----
def section_macro(doc, detail=False):
    h1(doc, "二、宏观维度 · 社零 / 供给 / 房地产")
    para(doc, "本章数据主体来自国家统计局2026-07-15发布的6月及上半年经济数据，交叉中指研究院、克而瑞房地产口径。"
              "家装建材家具行业作为地产后周期消费，需求端与竣工/二手房交易高度相关。", indent=True)

    h2(doc, "2.1　需求端 · 限额以上单位零售额同比（6月单月）")
    rs = [x for x in macro["macro"]["retailSales"] if "6月单月" in x["metric"] or "社零" in x["metric"]]
    rows = [[x["metric"].replace("限额以上单位","").replace("（6月单月）",""),
             x["value"], x["yoy"], x.get("publishDate","")] for x in rs]
    make_table(doc, ["指标","数值","同比","发布日"], rows, widths=[6.5,3.5,3.0,3.0])
    add_source_line(doc, "来源：国家统计局 https://www.stats.gov.cn/sj/zxfb/202607/t20260715_1964127.html")
    desc(doc, "家具类（-6.6%）、建材类（-10.5%）、家电类（-8.7%）三大家装相关品类6月单月全部深跌，"
              "且降幅显著大于社零大盘（+1.0%），其中建材类跌幅最深、与地产开工竣工负增长强相关。"
              "结构上，家装相关品类跑输大盘约8-11个百分点，说明当前疲弱是行业性、而非普遍消费收缩，"
              "运营动作应聚焦守住份额而非期待大盘回暖。")

    h2(doc, "2.2　供给端 · 建材工业与产量")
    ws = [x for x in macro["macro"]["wholesale"] if x["value"] != "待核实"]
    rows = [[x["metric"], x["value"], x["yoy"], x.get("publishDate","")] for x in ws]
    make_table(doc, ["指标","数值","同比","发布日"], rows, widths=[6.5,3.5,3.0,3.0])
    add_source_line(doc, "来源：国家统计局 6月工业增加值发布")
    desc(doc, "供给端呈结构性分化：非金属矿物制品业（-2.3%）、水泥（-5.6%）、平板玻璃（-5.3%）等地产强相关建材持续负增长，"
              "而金属制品业（含五金/金属家具，+7.8%）逆势正增长。这一分化提示：五金电工赛道景气度明显好于传统建材，"
              "天猫招商与资源投放可对五金/金属家具类目适度倾斜。（注：家具制造业分行业工业增加值国家统计局未单独公开，标待核实）")

    h2(doc, "2.3　房地产维度 · 销售 / 投资 / 开工 / 竣工（1-6月累计）")
    re_rows = []
    for x in (macro["realEstate"]["sales"][:3] + macro["realEstate"]["investment"] + macro["realEstate"]["newStart"]):
        re_rows.append([x["metric"], x["value"], x["yoy"], x.get("publishDate","")])
    make_table(doc, ["指标","数值","同比","发布日"], re_rows, widths=[6.5,3.8,2.7,3.0])
    add_source_line(doc, "来源：国家统计局 https://www.stats.gov.cn/sj/zxfb/202607/t20260715_1964126.html")
    desc(doc, "地产先行指标全面深跌：开发投资-18.0%、新开工-23.4%、竣工-23.7%。其中竣工端-23.7%是压制精装家居、"
              "软装、建材工程需求最关键的负向变量——竣工直接决定下游家居家装的入场时点。新开工深跌则预示2027-2028年"
              "需求将进一步承压。对天猫而言，这意味着家装建材行业未来1-2年缺乏总量红利，竞争将转为存量博弈。")

    if detail:
        h2(doc, "2.4　价格信号 · 70城房价与百城均价")
        pr = macro["realEstate"]["pricing"]
        rows = [[x["metric"], x["value"], x["yoy"]] for x in pr]
        make_table(doc, ["指标","数值/结构","同比方向"], rows, widths=[6.0,6.5,3.5])
        add_source_line(doc, "来源：国家统计局70城 + 中指研究院百城")
        desc(doc, "6月70城新房环比上涨面扩至20城，一线（上海/杭州/深圳/广州）温和企稳，但二手房70城全部同比下降。"
                  "由于二手房是二次装修（局改/翻新）需求主力，二手房价格持续走弱压制了改善型装修意愿；"
                  "而新房结构性企稳主要由核心城市高端项目拉动，对应的是高端全屋定制/精装配套需求，呈两极分化。")

        h2(doc, "2.5　本章数据缺口与冲突")
        for g in macro.get("gaps", [])[:5]:
            bullet(doc, g, size=10)
        para(doc, "口径冲突（双列保留，供人工判断）：", bold=True, indent=False, size=10.5)
        for c in macro.get("conflicts", [])[:3]:
            bullet(doc, c if isinstance(c,str) else json.dumps(c,ensure_ascii=False), size=10)
    page_break(doc)

# ---- doudian compass ----
def section_doudian(doc, detail=False):
    h1(doc, "三、抖音家装建材大盘 · 抖店罗盘（近7天 07/09-07/15）")
    para(doc, f"数据来源：抖店罗盘·类目概览（{dd['sourceUrl']}），经陪跑扫码登录后浏览器自动化抽取，采集时间 {dd['scrapedAt']}。"
              f"{dd['note']}", indent=True)

    h2(doc, "3.1　核心大盘指标（较上周期）")
    cm = dd["coreMetrics"]
    label = {"payGmv":"支付GMV","itemsSold":"销量件数","buyers":"支付买家数","orders":"订单数",
             "onlineSku":"在线商品数","activeSku":"动销商品数","pricePerItem":"件单价","arpu":"客单价"}
    rows = [[label[k], v["range"], v["wow"]] for k,v in cm.items()]
    make_table(doc, ["指标","数量级区间","较上周期(WoW)"], rows, widths=[4.5,6.0,4.0])
    add_source_line(doc, "来源：抖店罗盘·类目概览 智能家居/家装建材")
    desc(doc, "抖音家装建材大盘7天支付GMV ¥4.5-5亿、环比+1.51%，温和增长。订单数(+7.00%)与买家数(+4.89%)增速高于GMV，"
              "而件单价(+6.91%)、客单价(+6.10%)同步上行——说明本期增长由客单价拉动、并非纯流量红利，"
              "用户在抖音购买家装建材的件均金额在抬升，中高价商品跑量能力增强。")

    h2(doc, "3.2　子类目GMV排行（TOP12）")
    subs = dd["subCategories"][:12]
    rows = [[s["rank"], s["name"], s["gmvRange"], s["share"], s["unitPrice"], s["arpu"]] for s in subs]
    make_table(doc, ["排名","子类目","GMV区间","占比","件单价","客单价"], rows, widths=[1.3,3.3,3.4,2.4,2.3,2.3])
    add_source_line(doc, "来源：抖店罗盘 31个子类目（此处列TOP12）")
    desc(doc, "全屋智能（25-30%）+ 卫浴建材（20-25%）双龙头合计约占大盘一半。关键差异在件单价：全屋智能¥200-300 vs "
              "卫浴建材¥100-200——智能硬件在抖音已建立高价心智并跑量，卫浴建材走量为主。第3-11名（卫浴五金/水槽龙头/"
              "胶粘剂/防护材料/油漆/涂料等）多为¥10-50低件单价的耗材属性，呈长尾格局。对天猫而言，全屋智能是"
              "抖音已验证可高价跑量的赛道，与用户跟进的AI锁/AI摄像头方向一致，应重点做跨平台竞品对标。")

    h2(doc, "3.3　价格带分布")
    pb = dd["priceBands"]
    rows = [[b["rank"], b["band"], b["gmvRange"], b["share"]] for b in pb]
    make_table(doc, ["排名","价格带(元)","GMV区间","GMV占比"], rows, widths=[1.6,3.4,5.0,4.0])
    add_source_line(doc, "来源：抖店罗盘 5价格带")
    desc(doc, "价格带极度头重脚轻：¥39以上单档占75-80%GMV，¥19-39占10-15%，¥19以下三档合计不足15%。"
              "这颠覆了『抖音=低价白牌』的旧认知——抖音家装建材已由中高端主导，¥39以上是绝对主力战场。"
              "天猫在与抖音竞争时，不应简单以低价应对，而应在¥39以上区间比拼商品力、内容与信任心智。")

    h2(doc, "3.4　成交载体结构")
    cr = dd["carriers"]
    rows = [[c["name"], f"{c['share']}%", c["gmvRange"]] for c in cr]
    make_table(doc, ["成交载体","GMV占比","GMV区间"], rows, widths=[4.5,3.5,6.0])
    add_source_line(doc, "来源：抖店罗盘 5载体")
    desc(doc, "直播（54.80%）+ 商品卡（21.87%）+ 短视频（19.70%）合计96.4%，图文仅0.10%。这说明图文种草模型在抖音"
              "家装建材类目基本失效，直播与货架（商品卡）是主战场。值得注意的是商品卡占比超两成，反映抖音"
              "『货架电商』在家装建材已成型，用户主动搜索/复购行为显著，并非纯冲动直播下单。")

    if detail:
        h2(doc, "3.5　流量渠道来源（9渠道）")
        ch = dd["channels"]
        rows = [[c["name"], c["gmvRange"], c["share"], c["unitPrice"], c["arpu"]] for c in ch]
        make_table(doc, ["渠道","GMV区间","占比","件单价","客单价"], rows, widths=[3.6,3.4,2.6,2.4,2.4])
        add_source_line(doc, "来源：抖店罗盘 9渠道")
        desc(doc, "渠道TOP3为短视频引流（25-30%）+直播推荐（20-25%）+搜索（15-20%），合计60-70%。搜索占比达15-20%不低，"
                  "印证『抖音搜家装』心智已成型；个人主页&店铺&橱窗、关注来源等私域渠道件单价/客单价（¥200-400）"
                  "明显高于公域，说明沉淀粉丝的复购价值高，品牌自播+私域运营是提客单的关键。")

        h2(doc, "3.6　TOP15热销单品（¥39以上主力价格带）")
        tp = dd["topProducts_39plusBand"]
        rows = [[p["rank"], p["title"][:34], p["gmvRange"]] for p in tp]
        make_table(doc, ["排名","商品标题（截断）","GMV区间"], rows, widths=[1.4,10.6,3.0], fsize=8.5)
        add_source_line(doc, "来源：抖店罗盘 TOP15单品（¥39以上价格带）")
        desc(doc, "TOP15单品中，智能锁占4席（掌静脉/掌静脉/双子星/联想T30pro）、监控摄像头占4席（步猎鹰/神眸/斯麦特/AUX奥克斯），"
                  "两类合计8席、过半。这与用户负责跟进的『AI智能锁』『AI监控摄像头』两大赛道高度重叠——"
                  "建议将上述具体品牌（步猎鹰/神眸/斯麦特/因硕/阿尔法极光等白牌新锐）立即纳入天猫竞品追踪清单，"
                  "监测其是否向天猫渗透及定价策略。")

    h2(doc, "3.7　抖音大盘核心结论")
    for ins in dd["insights"]:
        bullet(doc, ins, size=10.5)
    page_break(doc)

# ---- platform dynamics ----
def section_platform(doc, detail=False):
    h1(doc, "四、电商平台竞争动态")
    para(doc, f"本章汇总{WINDOW}内各主要电商平台在家装建材家具品类的战报、战略与政策动向，均来自公开可查链接。", indent=True)
    platmeta = {"jd":"京东","douyin":"抖音","pdd":"拼多多","xhs":"小红书",
                "tmall_taobao":"淘宝天猫","shipinhao":"视频号","kuaishou":"快手"}
    total = 0
    for k, name in platmeta.items():
        items = plat["platforms"].get(k, [])
        real = [it for it in items if it.get("type") != "待核实"]
        if not real:
            continue
        h2(doc, f"4.{list(platmeta).index(k)+1}　{name}（{len(real)}条）")
        rows = [[it["date"], it["title"][:40], it.get("category","通用")] for it in real]
        make_table(doc, ["日期","动态标题","子行业"], rows, widths=[2.2,10.5,2.5], fsize=9)
        total += len(real)
        if detail:
            for it in real:
                para(doc, it["title"], bold=True, size=10.5, indent=False)
                para(doc, it.get("summary",""), size=10.5)
                add_source_line(doc, f"来源：{it.get('source','')}　{it.get('url','')}")
        desc(doc, _plat_desc(k, real))
    # cross platform
    cp = plat.get("crossPlatform", [])
    if cp:
        h2(doc, f"4.{len(platmeta)+1}　跨平台趋势（{len(cp)}条）")
        rows = [[it["date"], it["title"][:40], it.get("category","通用")] for it in cp]
        make_table(doc, ["日期","趋势标题","子行业"], rows, widths=[2.2,10.5,2.5], fsize=9)
        desc(doc, "跨平台看，618后家居业整体转向『价值重构』：京东整装翻倍、以旧换新10%补贴、下沉与老房翻新成新增长极，"
                  "低价高量模式被头部商家明确否定。这一信号对天猫的启示是：家装建材竞争正从价格战转向服务/整装/以旧换新的价值战。")
    page_break(doc)

def _plat_desc(k, items):
    if k=="jd":
        return ("京东是本期最активный平台：家装整装GMV+140%高增、618送装一体订单+120%、国补第三批625亿续力、"
                "京东MALL线下扩张至30店。京东以『整装+局改+安心住服务体系+线下MALL』构建全链路，"
                "对天猫家装形成正面竞争，尤其在整装与大件送装环节抢跑明显。")
    if k=="douyin":
        return ("抖音端一面是618家居家装战报（喜临门床垫13.5亿、九牧卫浴榜首、爱果乐多榜TOP1），"
                "一面是收紧商家资质合规。内容电商仍是家居品牌种草与直播成交的核心增量场，天猫需关注品牌预算向抖音的分流。")
    if k=="pdd":
        return ("拼多多披露以家电大类为主（洗地机+1100%、洗烘一体+700%），家装家具单独战报未公开；"
                "同期因『百亿补贴虚标』被北京市监局约谈。提示低价平台的补贴打法正受监管约束。")
    if k=="xhs":
        return ("小红书作为家居『内容驱动决策』核心阵地（爱果乐曝光6000w+、人体工学椅阅读渗透+248%），"
                "但首次被市监局纳入618约谈名单。小红书种草-成交链路对天猫家装的引流价值需持续跟踪。")
    if k=="tmall_taobao":
        return ("天猫618家装家居表现稳健：源氏木语/林氏/九牧领跑家装榜，慕思领跑家居榜，睡眠经济与轻定制突出、"
                "40000+品牌成交翻倍。作为主场，天猫在品牌矩阵与总裁直播上仍具规模优势，但需警惕京东整装与抖音内容的两翼夹击。")
    return "本平台在窗口期内有公开动态，详见上表来源链接。"

# ---- policy ----
def section_policy(doc, detail=False):
    h1(doc, "五、政策 · 标准 · 行业动态")
    para(doc, f"本章覆盖{WINDOW}内影响家装建材家具行业的政策法规、国家/团体标准及行业协会动态。", indent=True)

    h2(doc, "5.1　政策与标准")
    pol = policy["policy"]
    rows = [[p["issueDate"], p["title"][:42], p.get("category",""), "/".join(p.get("subIndustry",[]))[:16]] for p in pol]
    make_table(doc, ["发布日","政策/标准","类别","涉及子行业"], rows, widths=[2.2,7.5,2.3,3.2], fsize=9)
    add_source_line(doc, "来源：国家市场监管总局/发改委/商务部/各标准委 + 权威媒体")
    if detail:
        for p in pol:
            para(doc, p["title"], bold=True, size=10.5, indent=False)
            para(doc, p["summary"], size=10.5)
            para(doc, "影响：" + p.get("impact",""), size=10.5, color=RGBColor(0x8a,0x63,0x00))
            add_source_line(doc, f"来源：{p.get('source','')}　{p.get('url','')}")
    desc(doc, "本期政策主线是『提门槛、促出清』：GB18580-2025甲醛新国标（E0级强制）、智能卫浴可靠性国标、住宅全案整装服务标准"
              "三项标准同步落地，环保与质量门槛全面抬升，利好头部品牌（欧派/索菲亚/兔宝宝/九牧/恒洁）先发受益、"
              "加速中小厂出清。国补层面延续以旧换新（家装厨卫15%），并有湖南等地新增智能家居品类，构成需求端为数不多的正向拉动。"
              "天猫招商应向合规头部倾斜，并用足国补品类的流量红利。")

    h2(doc, "5.2　行业动态（协会/展会/景气指数）")
    ind = policy["industry"]
    rows = [[x["date"], x["title"][:42], x.get("topic","")] for x in ind]
    make_table(doc, ["日期","行业动态","类型"], rows, widths=[2.2,10.0,3.0], fsize=9)
    add_source_line(doc, "来源：CBMF/CBMCA/CBDA + 新华网/新京报等")
    desc(doc, "行业景气持续走弱：1-5月建材行业营收-15.3%、利润-68.6%，6月BHI环比-16.34进入淡季；"
              "广州建博会（50万㎡、近2000家）聚焦无醛整家/适老化/绿色低碳/AI，反映供给侧在存量竞争中主动向"
              "『健康+智能+服务』升级。景气数据印证宏观判断，展会方向则指明产品升级路径。")
    page_break(doc)

# ---- merchant ----
def section_merchant(doc, detail=False):
    h1(doc, "六、头部商家动态")
    para(doc, f"本章汇总{WINDOW}内家装建材家具头部上市公司公告与品牌大事件。", indent=True)
    mer = policy["merchant"]
    rows = [[m["date"], m["brand"], m.get("type",""), m["summary"][:30]+"…"] for m in mer]
    make_table(doc, ["日期","品牌","类型","摘要"], rows, widths=[2.0,3.2,2.6,7.5], fsize=9)
    add_source_line(doc, "来源：上证报/证券时报/东方财富/同花顺 + 各公司公告")
    if detail:
        for m in mer:
            para(doc, f"{m['brand']}（{m.get('type','')}｜{m['date']}）", bold=True, size=10.5, indent=False)
            para(doc, m["summary"], size=10.5)
            add_source_line(doc, f"来源：{m.get('source','')}　{m.get('url','')}")
    desc(doc, "头部商家动态呈现『定制承压、卖场与建材分化』格局：定制家居三巨头（欧派预降50-60%、志邦由盈转亏、"
              "索菲亚预降78-85%）业绩集体探底，是地产后周期与竞争加剧的直接体现；红星美凯龙扭亏为盈反映卖场端出租率企稳；"
              "九牧主导智能卫浴国标、兔宝宝获国家科技奖则显示头部通过标准与技术构筑壁垒。对天猫运营的含义："
              "定制家具赛道整体承压将传导至平台GMV，需在卫浴（国标+国补双驱动）、智能硬件等高景气赛道寻找结构性增量。")
    page_break(doc)

# ---- source depth report ----
def section_sources(doc):
    h1(doc, "七、附录：数据源采集深度报告")
    para(doc, srcidx["note"], indent=True)
    para(doc, "采集深度分级：0=未取到/阻塞，1=浅层（仅标题/目录），2=中层（关键指标/结构化摘要），3=深层（完整明细/多页/可下载表）。",
         indent=True, size=10)
    rows = []
    for s in srcidx["sources"]:
        rows.append([s["layer"], s["name"][:24], s["login"][:8], str(s["depth"]),
                     str(s.get("count",""))[:14], s.get("blocker","")[:30]])
    make_table(doc, ["层","数据源","登录","深度","采集量","阻塞/备注"], rows,
               widths=[1.0,4.0,1.8,1.2,2.8,4.5], fsize=8.5)
    sm = srcidx["summary"]
    desc(doc, f"共接入{sm['totalSources']}个数据源：深层采集(depth3){sm['depth3']}个、中层{sm['depth2']}个、浅层{sm['depth1']}个、"
              f"阻塞(depth0){sm['depth0_blocked']}个。{sm['blockedNote']} "
              f"公开层已采集{sm['publicItemsCollected']}。5个强登录源阻塞原因分别为：千瓜家居行业为会员墙、"
              "京麦需POP商家账号、京准通需广告主账号、蝉妈妈待后台入口URL、微信文档被公司网络DLP拦截——"
              "均非采集能力问题，待用户提供有效账号/入口后可即时补采。")

    if True:
        h2(doc, "7.1　阻塞源明细与破解路径")
        for s in srcidx["sources"]:
            if s["depth"] == 0:
                para(doc, f"● {s['name']}（{s['layer']}层）", bold=True, size=10.5, indent=False)
                para(doc, f"阻塞原因：{s['blocker']}", size=10)
    page_break(doc)

# ---- board conclusion ----
def section_conclusion(doc):
    h1(doc, "八、总体结论与天猫运营建议")
    h2(doc, "8.1　本期核心判断")
    for i, ins in enumerate(macro.get("keyInsightsForBoard", []), 1):
        para(doc, f"{i}. {ins}", size=11, indent=False)
    h2(doc, "8.2　对天猫家装运营的行动建议")
    actions = [
        ("守份额而非等回暖", "上半年家具-6.6%/建材-10.5%大幅跑输社零，行业缺乏总量红利，KPI设定与资源投放应以守住类目份额、抢竞品份额为主基调，不宜押注大盘自然反弹。"),
        ("押注两条高景气赛道", "卫浴厨房（智能卫浴国标+国补双驱动）与全屋智能（抖音已验证¥200-300高价跑量、AI锁/AI摄像头8席TOP15）是结构性增量，建议在招商、流量、活动资源上重点倾斜。"),
        ("正视中高价主战场", "抖音家装建材¥39以上占75-80%GMV，证明中高端可跑量。天猫应在中高价区间比拼商品力/内容/信任，而非陷入低价内卷（拼多多百亿补贴已被约束）。"),
        ("跨平台竞品对标", "将抖店罗盘TOP15白牌新锐（步猎鹰/神眸/斯麦特/因硕/阿尔法极光等）纳入天猫竞品追踪，监测其向天猫的渗透与定价。"),
        ("防守整装抢跑", "京东整装GMV+140%、局改+安心住服务体系抢跑，天猫需关注整装/局改/以旧换新/送装一体等价值链环节的竞品动向与自身补位。"),
        ("向合规头部倾斜", "三项标准（甲醛/卫浴可靠性/整装服务）落地加速中小出清，招商与流量分配应向合规头部品牌倾斜，规避标准合规风险商家。"),
    ]
    for name, txt in actions:
        para(doc, f"● {name}：{txt}", size=11, indent=False)
    h2(doc, "8.3　下期补采计划（待用户授权）")
    for b in ["抖店罗盘剩余3页（榜单/趋势/竞争分析）——能力可达，下期补采",
              "京麦商智：待用户切换POP商家账号扫码陪跑",
              "京准通：待用户提供京东广告主账号授权",
              "蝉妈妈后台榜单：待用户提供其后台书签URL或陪跑截图",
              "千瓜家居行业大盘：需尊享版会员权限（会员墙），或改用其他小红书数据源",
              "小红书千帆/蒲公英量化流量：需对应商业产品账号权限"]:
        bullet(doc, b, size=10.5)

# =========================================================
#  BUILD BRIEF (汇报稿 8-12页)
# =========================================================
def build_brief(path):
    doc = Document()
    style_doc_defaults(doc)
    for s in doc.sections:
        s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)
    cover(doc, "汇报稿（精要版）")
    section_monthly(doc, detail=False)
    section_macro(doc, detail=False)
    section_doudian(doc, detail=False)
    section_platform(doc, detail=False)
    section_policy(doc, detail=False)
    section_merchant(doc, detail=False)
    section_conclusion(doc)
    doc.save(path)
    return path

# =========================================================
#  BUILD DETAIL (详报 20-30页)
# =========================================================
def build_detail(path):
    doc = Document()
    style_doc_defaults(doc)
    for s in doc.sections:
        s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)
    cover(doc, "详报（完整版）")
    # 摘要
    h1(doc, "报告摘要（Executive Summary）")
    para(doc, mh["monthlySummary"], indent=True)
    para(doc, "本报告分八章：本月要闻、宏观维度、抖音大盘（抖店罗盘）、电商平台竞争动态、政策与标准、"
              "头部商家动态、数据源采集深度报告、总体结论与运营建议。每张数据表后均附自然语言解读，"
              "所有数据标注来源与发布日期，不可核实项标『待核实』，冲突信息双列保留。", indent=True)
    page_break(doc)
    section_monthly(doc, detail=True)
    section_macro(doc, detail=True)
    section_doudian(doc, detail=True)
    section_platform(doc, detail=True)
    section_policy(doc, detail=True)
    section_merchant(doc, detail=True)
    section_sources(doc)
    section_conclusion(doc)
    doc.save(path)
    return path

if __name__ == "__main__":
    outdir = sys.argv[1] if len(sys.argv) > 1 else ROOT
    os.makedirs(outdir, exist_ok=True)
    b = build_brief(os.path.join(outdir, "家装建材家具行业竞争情报月报_汇报稿.docx"))
    print("BRIEF ->", b)
    d = build_detail(os.path.join(outdir, "家装建材家具行业竞争情报月报_详报.docx"))
    print("DETAIL ->", d)
